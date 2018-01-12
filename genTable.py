from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, Color, Alignment

year = 2018

# create word table

document = Document()

# styles = document.styles
# for s in styles:
# 	if s.type == WD_STYLE_TYPE.TABLE:
# 			document.add_paragraph("Table style is : " + s.name)
# 			document.add_table(3, 3, style = s)
# 			document.add_paragraph("\n")

# if (year % 4 == 0) & (year % 100 != 0):
# 	year_long = 366
# elif year % 400 == 0:
# 	year_long = 366
# else:
# 	year_long = 365

#year_end = s + datetime.timedelta(days = year_long - 1)

s = datetime.datetime(year, 1, 1)
year_end = datetime.datetime(year, 12, 31)
weeks = year_end.strftime("%W")

delta = datetime.timedelta(days = 6 - s.weekday())

table = document.add_table(rows = int(weeks), cols = 5, style = "Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER

heading_cells = table.rows[0].cells
heading_cells[0].text = str(year)
heading_cells[1].text = '书'
heading_cells[2].text = '菜'
heading_cells[3].text = '电影'
heading_cells[4].text = '算法'

first_col = table.columns[0].cells

for i in range(1, len(table.rows)):
	e = "(" + s.strftime("%m/%d") + '--' + (s + delta).strftime("%m/%d") + ")"
	first_col[i].text = "第%d周\n" % i + e
	s = s + datetime.timedelta(days = 7)
	delta = datetime.timedelta(days = 6)

remain = (year_end - (s - datetime.timedelta(days = 1))).days
if (remain >= 5):
	row = table.add_row()
	e = "(" + s.strftime("%m/%d") + '--' + year_end.strftime("%m/%d") + ")"
	row.cells[0].text = "第53周\n" + e

document.add_page_break()
document.save('demo.docx')

# create xlsx table

x_document = Workbook()
sheet = x_document.create_sheet()
sheet.title = "年度报表"

sheet['A1'] = str(year)
sheet['A1'].font = Font(bold = True, italic = True)
row0 = ['书', '电影', '算法']
b = Font(bold = True)
a = Alignment(horizontal = "center", wrap_text = True)
for i in range(2, len(row0) + 2):
	sheet.cell(row = 1, column = i, value = row0[i-2])
	sheet.cell(row = 1, column = i).font = b
	sheet.cell(row = 1, column = i).alignment = a

s = datetime.datetime(year, 1, 1)
delta = datetime.timedelta(days = 6 - s.weekday())
for i in range(2, int(weeks) + 1):
	e = "(" + s.strftime("%m/%d") + '--' + (s + delta).strftime("%m/%d") + ")"
	sheet.cell(row = i, column = 1, value = "第%d周\n" % (i - 1) + e)
	sheet.cell(row = i, column = 1).font = b
	sheet.cell(row = i, column = 1).alignment = a
	s = s + datetime.timedelta(days = 7)
	delta = datetime.timedelta(days = 6)

remain = (year_end - (s - datetime.timedelta(days = 1))).days
if (remain >= 5):
	e = "(" + s.strftime("%m/%d") + '--' + year_end.strftime("%m/%d") + ")"
	sheet.cell(row = 54, column = 1, value = "第53周\n" + e)
	sheet.cell(row= 54, column = 1).font = b
	sheet.cell(row = 54, column = 1).alignment = a

x_document.save('demo.xlsx')
